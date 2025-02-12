from docx import Document as dok
from copy import deepcopy
import subprocess
import os
from typing import List


def subsection_with_explanation(obj: any, Heading: str, Description: str) -> None:
    """
    Creates a subsection with a heading and description in LaTeX format.

    :param obj: The LaTeX document object to which the content is written.
    :type obj: any
    :param Heading: The title of the section/subsection.
    :type Heading: str
    :param Description: The description text for the section/subsection.
    :type Description: str
    """
    obj.write("\section*{" + escape_special_chars(Heading) + "}\n")
    obj.write(escape_special_chars(Description) + "\n")


def table_description(obj: any, table_list: List[str], table_para_list: List[str], table_col_para_list: List[str]) -> None:
    """
    Creates a description for a table and its sub-elements in LaTeX format.

    :param obj: The LaTeX document object to which the content is written.
    :type obj: any
    :param table_list: List of table header data.
    :type table_list: List[str]
    :param table_para_list: List of table paragraph data.
    :type table_para_list: List[str]
    :param table_col_para_list: List of table column data.
    :type table_col_para_list: List[str]
    """
    k = 0
    obj.write("\\begin{description}\n")
    while k < len(table_list):
        if (
            table_list[len(table_list) - 1]
            and "=" in table_list[k + 1]
            and ":" in table_list[k]
        ):
            obj.write(
                "\\item[\\small "
                + escape_special_chars(table_heading(table_col_para_list, k))
                + "]\mbox{}\n"
            )
            obj.write(
                escape_special_chars(table_subdescription(table_col_para_list, k))
                + "\\"
                + "\\"
                + "\n"
            )
            new_list = itemize_function(table_para_list)
            for para_item in new_list:
                obj.write(para_item)
        else:
            obj.write(
                "\\item[\\small "
                + escape_special_chars(table_list[k])
                + "]"
                + escape_special_chars(table_list[k + 1])
                + "\n"
            )
        k += 2
    obj.write("\\end{description}" + "\n")


def itemize_function(table_paragraph_list: List[str]) -> List[str]:
    """
    Converts a list of table paragraphs into an itemized list in LaTeX format.

    :param table_paragraph_list: List of table paragraphs.
    :type table_paragraph_list: List[str]
    :return: A list of LaTeX itemized list strings.
    :rtype: List[str]
    """
    print(table_paragraph_list)
    itemize_list = []
    itemize_list.append("\\begin{itemize}\n")
    while len(table_paragraph_list) > 0:
        if "=" in table_paragraph_list[0]:
            item_str = (
                "\t" + "\\item " + escape_special_chars(table_paragraph_list[0]) + "\n"
            )
            itemize_list.append(item_str)
            table_paragraph_list.pop(0)
            print('"=" read: ')
            print(table_paragraph_list)
            if (len(table_paragraph_list) <= 0) or (":" in table_paragraph_list[0]):
                print("loop broken!")
                print(table_paragraph_list)
                break
        else:
            table_paragraph_list.pop(0)

    itemize_list.append("\\end{itemize}" + "\\" + "\n")
    return itemize_list


def table_heading(col_1_para_list: List[str], num: int) -> str:
    """
    Returns the heading for a table from the first column.

    :param col_1_para_list: List of paragraph data from the first column.
    :type col_1_para_list: List[str]
    :param num: Index to access the correct heading.
    :type num: int
    :return: The heading of the table.
    :rtype: str
    """
    return col_1_para_list[num]


def table_subdescription(col_1_para_list: List[str], num: int) -> str:
    """
    Returns the subdescription for a table from the first column.

    :param col_1_para_list: List of paragraph data from the first column.
    :type col_1_para_list: List[str]
    :param num: Index to access the correct subdescription.
    :type num: int
    :return: The subdescription of the table.
    :rtype: str
    """
    return col_1_para_list[num + 1]


def escape_special_chars(textWUnd: str) -> str:
    """
    Escapes special LaTeX characters in a given string.

    :param textWUnd: The string containing potential LaTeX special characters.
    :type textWUnd: str
    :return: The string with LaTeX special characters escaped.
    :rtype: str
    """
    new_str = ""
    for i in range(len(textWUnd)):
        if textWUnd[i] == "_":
            handle_char = "\\_ "
            new_str = new_str + handle_char
        elif textWUnd[i] == "&":
            handle_char = "\\& "
            new_str = new_str + handle_char
        elif textWUnd[i] == "$":
            handle_char = "\\$ "
            new_str = new_str + handle_char
        elif textWUnd[i] == "%":
            handle_char = "\\% "
            new_str = new_str + handle_char
        elif textWUnd[i] == "{":
            handle_char = "\\{ "
            new_str = new_str + handle_char
        elif textWUnd[i] == "}":
            handle_char = "\\} "
            new_str = new_str + handle_char
        elif textWUnd[i] == "#":
            handle_char = "\\# "
            new_str = new_str + handle_char
        elif textWUnd[i] == "~":
            handle_char = "\\textasciitilde "
            new_str = new_str + handle_char
        elif textWUnd[i] == "^":
            handle_char = "\\textasciicircum "
            new_str = new_str + handle_char
        elif textWUnd[i] == "\\":
            handle_char = "\\textbackslash "
            new_str = new_str + handle_char
        else:
            new_str = new_str + textWUnd[i]
    return new_str


def clean_list(table_list: List[str]) -> List[str]:
    """
    Cleans a list by removing empty and whitespace strings.

    :param table_list: List of table elements.
    :type table_list: List[str]
    :return: Cleaned list without empty or whitespace strings.
    :rtype: List[str]
    """
    cl_list = []
    for elem in table_list:
        if elem != "" and elem != " ":
            cl_list.append(elem)
    return cl_list


def Insert_description_file(tex_file_without_description: str) -> str:
    """
    Inserts description from a Word document into a LaTeX file.

    - Creates a copy of the Word template.
    - Opens the document to collect paragraph and table data.
    - Generates a LaTeX description file.
    - Inserts the description into a LaTeX document.

    :param tex_file_without_description: Path to the LaTeX file without the description.
    :type tex_file_without_description: str
    :return: Path to the LaTeX file with the description inserted.
    :rtype: str
    """
    # Create new file with study number name from standard template file and open it for the user to make changes
    sample_doc = dok(r"path_to_project\DESCPR_files\std_word_template\Standard_Latex_Description_File.docx")
    description_folder_path = r"path_to_project\DESCPR_files"
    copy_of_content = deepcopy(sample_doc)

    new_doc_file_name = os.path.join(description_folder_path, tex_file_without_description.split("\\")[-1].replace(".tex", "_description.docx"))
    copy_of_content.save(new_doc_file_name)

    proc_doc = subprocess.Popen(f"start {new_doc_file_name} && pause", shell=True)
    proc_doc.wait()
    prompt_value = input("Have you setup the word description file? (y/n): ")

    while True:
        if prompt_value == "y" or prompt_value == "Y":
            break
        elif prompt_value == "n" or prompt_value == "N":
            print("User exited.")
            exit()
        prompt_value = input("Have you setup the word description file? (y/n): ")

    document = dok(new_doc_file_name)
    doc_tex_filename = os.path.join(description_folder_path, tex_file_without_description.split("\\")[-1].replace(".tex", "_description.tex"))
    # Collecting all paragraph objects
    all_paragraphs = document.paragraphs
    print("\n\nNo of paragraphs read: " + str(len(all_paragraphs)))

    # Empty List for collecting each paragraph text generated from paragraphs objects
    names_explanation_list = []

    for num, para in enumerate(all_paragraphs, start=1):
        if para.text == "":
            continue
        print(str(num) + "-" + para.text)
        names_explanation_list.append(para.text)

    # Reads text from each cell in each Table
    all_tables = document.tables

    # Create a dictionary for generating empty lists
    #  Indexing according to number of table objects stored in 'all_tables' variable
    dict_table = {}

    # Assigning read tables as dict value object to its respective key names in the dictionary
    for obj_num, obj in enumerate(all_tables, start=1):
        list_for_table = "table_" + str(obj_num) + "_list"

        dict_table[list_for_table] = []
        if obj_num == len(all_tables):
            list_for_para = "table_" + str(obj_num) + "_para_list"
            list_for_col = "table_" + str(obj_num) + "_col_1_para_list"
            dict_table[list_for_para] = []
            dict_table[list_for_col] = []

    print("\n\nThe dictionary of table names contains the following elements")
    print(dict_table)
    print("\n\nNo of tables read: " + str(len(all_tables)))

    # Accessing text from each cell in each table
    for table_num, table in enumerate(all_tables, start=1):
        for row in table.rows:
            for cell in row.cells:
                dict_table["table_" + str(table_num) + "_list"].append(cell.text)

    for table_num, table in enumerate(all_tables, start=1):
        if table_num == len(all_tables):
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        dict_table["table_" + str(table_num) + "_para_list"].append(para.text)
                        if ":" in para.text or "=" not in para.text:
                            dict_table["table_" + str(table_num) + "_col_1_para_list"].append(para.text)

    ### Viewing list elements
    print(names_explanation_list)
    print(len(names_explanation_list))

    for key, value in dict_table.items():
        print("\n\n" + "List name:" + key, end="\n")
        print(value)

    key_list = list(dict_table.keys())
    print("\n")
    print("This is the key list")
    print(key_list)

    list_of_tables = key_list[0 : (len(key_list) - 2)]
    table_X_para_list = key_list[-2]
    table_X_col_1_para_list = key_list[-1]

    # Removes spaces and null strings from the list
    clean_table_X_col_1_para_list = clean_list(dict_table[table_X_col_1_para_list])
    print()
    print(clean_table_X_col_1_para_list)

    # Removing title due to its non-iterative and standard nature
    title = names_explanation_list.pop(0)

    # List for separating headings and descriptions (belonging to tables only)
    table_explanation_list = names_explanation_list[
        (len(names_explanation_list) - (len(all_tables) * 2)) :
    ]
    print("\n")
    print("The items for the table explanation are: ")
    print(table_explanation_list)
    print("\n")
    print("The items for the names explanation list are: ")
    print(names_explanation_list)

    # Creating a new tex document and populating it with corresponding objects
    with open(doc_tex_filename, "w") as TxF:
        TxF.write("\\section*{\\huge " + escape_special_chars(title) + "}" + "\n")
        z = 0

        while z < len(names_explanation_list):
            if names_explanation_list[z] in table_explanation_list:
                break
            TxF.write("\\section*{" + escape_special_chars(names_explanation_list[z]) + "}\n")
            TxF.write(escape_special_chars(names_explanation_list[z + 1]) + "\n")
            z += 2
        i = 0
        j = 0

        while i < len(table_explanation_list) and j < len(list_of_tables):
            print("i starts with:" + str(i))
            print("j starts with:" + str(j))

            subsection_with_explanation(TxF, table_explanation_list[i], table_explanation_list[i + 1])

            table_description(TxF, dict_table[list_of_tables[j]], 
                                        dict_table[table_X_para_list], clean_table_X_col_1_para_list)
            j += 1
            i += 2
            print("j ends with:" + str(j))
            print("i ends with:" + str(i))

    with open(tex_file_without_description) as ftex:
        lines_gen = iter(ftex.readlines())
        print(lines_gen)
        tex_temp = []
        while True:
            try:
                l = next(lines_gen)
                if "\\begin{document}" in l:
                    tex_temp.append(l)
                    
                    # input syntax where description file is appended
                    description_filepath_to_insert = doc_tex_filename.replace("\\", "/")
                    tex_temp.append(r"\input{" + f"{description_filepath_to_insert}" + "}\n" + r"\newpage")
                else:
                    tex_temp.append(l)
            except StopIteration:
                break

    new_tex_gen = iter(tex_temp)

    old_latex_file_with_description = f"{tex_file_without_description[:-4]}" + "_final.tex"

    with open(old_latex_file_with_description, "w") as f_new:
        while True:
            try:
                f_new.write(next(new_tex_gen))
            except StopIteration:
                break

    
    return old_latex_file_with_description
